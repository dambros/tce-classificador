# -*- coding: utf-8 -*-
import csv
import glob
import json
import logging.config
import os
import re
from datetime import datetime
from itertools import islice

from dateutil.parser import parse
from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from lxml import etree

import config


def main():
    results = []
    for docs_path in config.SHARED_FOLDER_PATHS:
        results.append(classify(docs_path))

    for result in results:
        write_xml(result)

    if config.SUMMARY_PATH:
        keys = ['path', 'parent', 'child']
        write_csv(config.SUMMARY_PATH, results, keys)


def classify(docs_path):
    path = get_path(docs_path, 'docx')
    last_mod_dates = get_last_modifications_dates()
    doc_list = glob.glob(path, recursive=True)

    results = []
    if last_mod_dates and docs_path in last_mod_dates:
        last_mod_date = parse(last_mod_dates[docs_path])
        filter_files_after(doc_list, last_mod_date)

    if doc_list:
        for doc in doc_list:
            result = {}
            try:
                root_category = get_root_category(doc)
                result['path'] = doc
                result['parent'] = root_category
                if root_category not in ['Geral', 'Despacho']:
                    sub_category = get_sub_category(doc, root_category)
                    if sub_category:
                        result['child'] = sub_category
                results.append(result)
            except PackageNotFoundError:
                logger.error(doc)

        save_current_modification_date(docs_path)

    return results


def write_xml(results):
    for result in results:
        xml_path = re.sub(r'^/', '', re.sub(r'docx$', 'xml', result['path']))
        xml_path = os.path.join(config.XML_FOLDER_PATH, xml_path)

        root = etree.Element('root')
        categories = etree.SubElement(root, 'categories')
        categories.text = result['parent']
        file_path = etree.SubElement(root, 'file_path')
        file_path.text = result['path']

        if 'child' in result.keys():
            categories.text += '/' + result['child']

        write_file(xml_path,
                   etree.tostring(root, pretty_print=True, encoding='unicode'))


def write_file(filename, content):
    create_directory_if_not_exists(filename)
    with open(filename, 'w') as file:
        file.write(content)


def write_csv(filename, content, keys=[]):
    create_directory_if_not_exists(filename)
    with open(filename, 'w') as file:
        writer = csv.DictWriter(file, keys, delimiter=';')
        writer.writeheader()
        for row in content:
            writer.writerows(row)


def create_directory_if_not_exists(filename):
    directory = os.path.dirname(filename)
    if not os.path.exists(directory):
        os.makedirs(directory)


def get_root_category(doc_path):
    document = Document(doc_path)
    paragraphs = document.paragraphs

    start_paragraphs = islice(paragraphs, config.FIRST_PAGE_PARAGRAPHS)
    for p in start_paragraphs:
        text = p.text.lower()
        if text:
            for name, values_list in config.ROOT_CATEGORIES.items():
                for value in values_list:
                    if re.search(r'\b' + value.lower() + r'\b', text.lower()):
                        return name

    reversed_paragraphs = reversed(paragraphs)
    last_paragraphs = islice(reversed_paragraphs, 20)
    for p in last_paragraphs:
        text = p.text.lower()
        if text:
            for name, values_list in config.ROOT_CATEGORIES.items():
                for value in reversed(values_list):
                    if re.search(r'\b' + value.lower() + r'\b', text.lower()):
                        return name

    return 'Geral'


def get_sub_category(doc_path, root_category):
    document = Document(doc_path)
    paragraphs = document.paragraphs

    start_paragraphs = islice(paragraphs, config.FIRST_PAGE_PARAGRAPHS)
    for p in start_paragraphs:
        text = p.text.lower()
        if text:
            for root, subs in config.SUB_CATEGORIES[root_category].items():
                for sub in subs:
                    if re.search(r'\b' + sub.lower() + r'\b', text.lower()):
                        return root


def get_last_modifications_dates():
    if os.path.isfile(config.DB_PATH):
        with open(config.DB_PATH, 'r') as f:
            return json.load(f)

    return None


def save_current_modification_date(root_folder):
    mod_dates = get_last_modifications_dates()
    if mod_dates:
        mod_dates[root_folder] = datetime.now().isoformat()
    else:
        mod_dates = {
            root_folder: datetime.now().isoformat()
        }

    write_file(config.DB_PATH,
               json.dumps(mod_dates, indent=1, ensure_ascii=False))


def filter_files_after(files, last_modified_date):
    for file in files[:]:
        file_modified_date = datetime.fromtimestamp(os.path.getmtime(file))
        if last_modified_date > file_modified_date:
            files.remove(file)


def get_path(folder, extension):
    path = f'{folder}/**/*.{extension}'
    return path


if __name__ == "__main__":
    logging.config.fileConfig(config.LOG_CONFIG_PATH)
    logger = logging.getLogger(os.path.basename(__file__))
    startTime = datetime.now()
    main()
    logger.info(f'Script finalizado em {datetime.now() - startTime}')
